from PyQt6.QtWidgets import (QDialog, QFormLayout, QLineEdit, QPushButton,
                             QMessageBox, QFileDialog, QDateEdit, QCheckBox)
from PyQt6.QtCore import QDate
from database import db
import os
from copy import deepcopy

try:
    from docx import Document
    from docx.table import Table
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ReportInfoDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("生成整改报告 - 信息录入")
        self.resize(400, 300)
        self.layout = QFormLayout(self)

        self.sys_name = QLineEdit()
        self.client_name = QLineEdit()
        self.report_org = QLineEdit()
        self.date_edit = QDateEdit()
        self.date_edit.setDate(QDate.currentDate())

        self.check_include_all = QCheckBox("包含所有测评项 (不仅仅是不符合项)")
        self.check_include_all.setChecked(False)

        self.btn_gen = QPushButton("开始生成")
        self.btn_gen.clicked.connect(self.accept)

        self.layout.addRow("系统名称:", self.sys_name)
        self.layout.addRow("客户单位名称:", self.client_name)
        self.layout.addRow("出具报告单位:", self.report_org)
        self.layout.addRow("报告日期:", self.date_edit)
        self.layout.addRow("", self.check_include_all)
        self.layout.addWidget(self.btn_gen)

    def get_info(self):
        return {
            "sys_name": self.sys_name.text(),
            "client_name": self.client_name.text(),
            "report_org": self.report_org.text(),
            "date": self.date_edit.date().toString("yyyy年MM月dd日"),
            "include_all": self.check_include_all.isChecked()
        }


class ReportGenerator:
    def __init__(self, parent_widget, project_id):
        self.parent = parent_widget
        self.project_id = project_id

    def generate(self):
        if not HAS_DOCX:
            QMessageBox.critical(self.parent, "错误", "未检测到 python-docx 库。\n请运行 pip install python-docx")
            return

        tpl_path = "模板.docx"
        if not os.path.exists(tpl_path):
            QMessageBox.critical(self.parent, "错误", "未找到 '模板.docx'。")
            return

        dialog = ReportInfoDialog(self.parent)
        if not dialog.exec(): return

        info = dialog.get_info()
        default_name = f"{info['sys_name']}整改报告.docx"
        save_path, _ = QFileDialog.getSaveFileName(self.parent, "保存报告", default_name, "Word Files (*.docx)")
        if not save_path: return

        try:
            doc = Document(tpl_path)

            # 1. 替换基础文本 (保留格式)
            self._replace_doc_text(doc, "[系统名称]", info['sys_name'])
            self._replace_doc_text(doc, "[客户单位名称]", info['client_name'])
            self._replace_doc_text(doc, "[出具报告单位名称]", info['report_org'])
            self._replace_doc_text(doc, "[报告生成日期]", info['date'])

            # 2. 定位锚点：找到包含 'point' 的表格，以此为基准进行插入
            proto_table_element = None
            anchor_index = -1
            parent_elm = None

            # 遍历寻找模板表格
            for i, block in enumerate(doc.element.body):
                if block.tag.endswith('tbl'):
                    # 这是一个表格，检查内容
                    tbl_text = ""
                    for text_elem in block.iter(qn('w:t')):
                        if text_elem.text: tbl_text += text_elem.text

                    if "point" in tbl_text:
                        # 找到了！
                        proto_table_element = deepcopy(block)
                        anchor_index = i
                        parent_elm = doc.element.body
                        break

            if proto_table_element is None:
                QMessageBox.warning(self.parent, "警告", "模板中未找到包含 'point' 的表格。")
                doc.save(save_path)
                return

            # 3. 彻底清理：删除模板中现有的【测评分类标题】和【模板表格】
            # 先删除表格
            parent_elm.remove(parent_elm[anchor_index])

            # 再尝试删除前面的 "[测评分类]" 段落 (如果存在)
            # 我们向上搜索最多10个元素
            for i in range(anchor_index - 1, max(-1, anchor_index - 10), -1):
                elm = parent_elm[i]
                if elm.tag.endswith('p'):
                    # 检查文字
                    txt = ""
                    for t in elm.iter(qn('w:t')):
                        txt += t.text
                    if "测评分类" in txt or "资产名称" in txt:
                        parent_elm.remove(elm)
                        # 删除后，索引减1
                        anchor_index -= 1

            # 4. 获取数据
            all_assessments = db.get_all_project_assessments(self.project_id)
            project_data = db.get_project_data(self.project_id)
            disabled_cats = project_data.get('disabled_categories', [])
            from database import STD_CATEGORIES

            # 按分类组织数据
            data_by_cat = {cat: [] for cat in STD_CATEGORIES}
            for record in all_assessments:
                cat = record.get('category')
                if cat in data_by_cat:
                    data_by_cat[cat].append(record)

            # 5. 原位插入逻辑 (In-Place Insertion)
            current_insert_idx = anchor_index
            total_items_generated = 0

            for category in STD_CATEGORIES:
                if category in disabled_cats: continue

                items = data_by_cat.get(category, [])
                if not items: continue

                # 筛选
                target_items = []
                for item in items:
                    status = item.get('status', '不适用')
                    if info['include_all'] or status in ['不符合', '部分符合']:
                        target_items.append(item)

                if not target_items: continue

                # A. 插入分类标题 (使用 doc.add_heading 生成并搬运)
                title_elm = self._generate_and_move_heading(doc, f"一、{category}", level=1)
                parent_elm.insert(current_insert_idx, title_elm)
                current_insert_idx += 1

                # 安全计算环境按资产分组
                if category == "安全计算环境":
                    assets_map = {}
                    for item in target_items:
                        a_name = item.get('asset_name', '未知资产')
                        if a_name not in assets_map: assets_map[a_name] = []
                        assets_map[a_name].append(item)

                    for asset_name, a_items in assets_map.items():
                        # 插入资产子标题
                        sub_elm = self._generate_and_move_heading(doc, f"资产: {asset_name}", level=2)
                        parent_elm.insert(current_insert_idx, sub_elm)
                        current_insert_idx += 1

                        for item in a_items:
                            # B. 插入表格
                            new_tbl = self._create_item_table(doc, proto_table_element, item, asset_name)
                            parent_elm.insert(current_insert_idx, new_tbl)
                            current_insert_idx += 1

                            # C. 插入标准间隔 (正常行距)
                            spacer = self._create_normal_spacer(doc)
                            parent_elm.insert(current_insert_idx, spacer)
                            current_insert_idx += 1

                            total_items_generated += 1
                else:
                    for item in target_items:
                        # B. 插入表格
                        new_tbl = self._create_item_table(doc, proto_table_element, item, category)
                        parent_elm.insert(current_insert_idx, new_tbl)
                        current_insert_idx += 1

                        # C. 插入标准间隔 (正常行距)
                        spacer = self._create_normal_spacer(doc)
                        parent_elm.insert(current_insert_idx, spacer)
                        current_insert_idx += 1

                        total_items_generated += 1

            doc.save(save_path)

            if total_items_generated == 0:
                QMessageBox.information(self.parent, "提示", "报告已生成，但为空。未发现不符合项。")
            else:
                QMessageBox.information(self.parent, "完成", f"报告生成成功！\n共生成 {total_items_generated} 个表格。")

        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self.parent, "错误", f"生成报告时出错:\n{str(e)}")

    def _generate_and_move_heading(self, doc, text, level=1):
        """
        利用 doc.add_heading 生成标准标题，然后将其从文档末尾分离并返回。
        这样可以确保正确应用了 'Heading 1' 等样式。
        """
        # 1. 在文档末尾生成标题
        p = doc.add_heading(text, level)
        # 2. 获取其 XML 元素
        p_element = p._element
        # 3. 从文档末尾移除 (因为我们稍后要手动 insert 到中间)
        # 注意：add_heading 会将其 append 到 body，所以我们从 body 移除最后添加的这个
        doc.element.body.remove(p_element)
        return p_element

    def _create_normal_spacer(self, doc):
        """
        创建一个标准高度的空行，用于分隔表格。
        不使用极小值，而是使用标准的段落间距。
        """
        # 1. 创建一个空段落
        p = doc.add_paragraph("")

        # 2. 设置为固定值行距，确保大概是一个空行的高度 (例如 12pt)
        # 如果直接用 add_paragraph("")，受 Normal 样式影响，可能会很大或很小
        # 这里强制设置一个适中的行距
        p_format = p.paragraph_format
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing = Pt(30)  # 12磅，约等于一行正文高度

        # 3. 同样的“搬运”逻辑
        p_element = p._element
        doc.element.body.remove(p_element)
        return p_element

    def _create_item_table(self, doc, proto_elm, item, asset_name):
        """复制并填充表格，返回 XML 元素"""
        tbl_copy = deepcopy(proto_elm)
        # 临时包装成 Table 对象以便处理
        new_table = Table(tbl_copy, doc)

        data_map = {
            "point": str(item.get('point') or ''),
            "requirement": str(item.get('requirement') or ''),
            "result": str(item.get('result') or ''),
            "zc": str(asset_name or ''),
            "risk": str(item.get('risk_level') or '中'),
            "suggestion": str(item.get('suggestion') or '')
        }

        self._fill_table_with_data(new_table, data_map)

        return tbl_copy

    def _replace_doc_text(self, doc, old_text, new_text):
        if not new_text: new_text = ""
        for p in doc.paragraphs:
            if old_text in p.text:
                for run in p.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

    def _fill_table_with_data(self, table, data_map):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:

                        # 1. 优先执行数据替换和颜色逻辑
                        for key, val in data_map.items():
                            if key in run.text:
                                run.text = run.text.replace(key, val)

                                # 如果是风险等级且为“高”，设置为红色
                                if key == "risk" and val == "高":
                                    run.font.color.rgb = RGBColor(255, 0, 0)

                        # 2. 统一设置字体为 华文仿宋 5号 (10.5pt)
                        run.font.name = '华文仿宋'
                        run.font.size = Pt(10.5)

                        # 设置中文字体支持
                        # 必须通过 element.rPr.rFonts 来设置东亚字体
                        r = run._element
                        if r.rPr is None:
                            r.get_or_add_rPr()
                        r.rPr.rFonts.set(qn('w:eastAsia'), '华文仿宋')